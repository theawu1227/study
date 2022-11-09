import pandas as pd
from openpyxl import load_workbook
import re
from docx import Document
from win32com import client as wc
import time


def read_dir(path):
    import os
    sheet_l = []
    for i in os.listdir(path):
        year = "".join(re.findall('\d{4}', i)) + '年'
        if year not in sheet_l:
            print(year)
            create_sheet(year)
            sheet_l.append(year)
        new_path = os.path.join(path, i)
        print(new_path)
        try:
            # doSaveAas(new_path,i)
            # read_docx(new_path, year)
            eord_table(new_path,year)
        except:
            continue
        time.sleep(1)
        # break


# 将doc 转成docx
def doSaveAas(new_path, i):
    word = wc.Dispatch('Word.Application')
    doc = word.Documents.Open(new_path)  # 目标路径下的文件
    copy_path = f"D:\桌面\\ruxuan\\{i}x"
    print(copy_path)
    doc.SaveAs(copy_path, 12, False, "", True, "", False, False, False, False)  # 转化后路径下的文件
    doc.Close()
    word.Quit()
    time.sleep(1)


def read_docx(file, year):
    doc = Document(file)  # 实例化
    # print(files)
    bool = 0
    school = ""
    for va in doc.paragraphs:
        if bool == 0:
            if (va.text).strip() != '附件':
                continue
            else:
                bool += 1
        # print('====')
        # print(va.text)
        if va.text == "":
            continue
        p = "".join(re.findall('\d+人', va.text))
        if p != "":
            school = re.sub('（(.*)', "", va.text)
        else:
            if school == "":
                continue
            else:
                # print(va.text)
                people = chuli(va.text)
                # print(people)
                people_l = people.split('  ')
                # print(people_l)
                first = ''
                for p in people_l:
                    if p == "":
                        continue
                    if len(p) == 1:
                        if first == "":
                            first = p
                            continue
                        else:
                            name = first + p
                            first = ""
                    else:
                        name = re.sub('（(.*)', "", p)
                    print(name)

                    #     name = re.sub('（(.*)', "", p)
                    #     print(school, name)
                    in_datas(school, name, year)


def run():
    year = input('year:')
    while True:
        schoolfile = input('school:')
        if schoolfile == "q":
            break
        try:
            l_num = "".join(re.findall('(\d+)人', schoolfile))
            l_num = int(l_num)
        except:
            print('学校不完整')
            schoolfile = input('school:')
            l_num = "".join(re.findall('(\d+)人', schoolfile))
            l_num = int(l_num)
        school = re.sub('（(.*)', "", schoolfile)
        print(school, l_num)
        peoples = 0
        while True:
            if peoples == l_num:
                break
            file = input('file:')
            people = chuli(file)
            # print(people)
            people_l = people.split('	')
            # print(people_l)
            first = ''
            for p in people_l:
                if p == "":
                    continue
                if len(p) == 1:
                    if first == "":
                        first = p
                        continue
                    else:
                        name = first + p
                        first = ""
                else:
                    name = re.sub('（(.*)', "", p)
                print(name)
                peoples += 1

                #     name = re.sub('（(.*)', "", p)
                #     print(school, name)
                try:
                    in_datas(school, name, year)
                except Exception as e:
                    print(e)


def chuli(test):
    n = re.findall('   ', test)

    if len(n) > 0:
        file = re.sub('   ', '  ', test)

        file = chuli(file)
    else:
        file = test
    return file


def in_datas(school, name, year):
    dic = {
        1: school,
        2: name,
    }
    wb = load_workbook('D:\\桌面\\东方学者.xlsx')
    ws = wb[year]
    ws.append(dic)
    wb.save('D:\\桌面\\东方学者.xlsx')


def create_sheet(sheet):
    wb = load_workbook('D:\\桌面\\东方学者.xlsx', )
    wb.create_sheet(sheet)
    wb.save('D:\\桌面\\东方学者.xlsx')


def create_table():
    dic = {
        '学校': [],
        '人名': []
    }
    df = pd.DataFrame(dic)
    df.to_excel('D:\\桌面\\东方学者.xlsx', index=False, sheet_name='2013年')


def eord_table(file,year):
    doc = Document(file)
    tables = doc.tables
    bool = 0
    school_l = []
    for va in doc.paragraphs:
        if bool == 0:
            if (va.text).strip() != '附件':
                continue
            else:
                bool += 1
        if (va.text).strip() not in ('附件',''):
            # print(va.text)
            school = re.sub('（(.*)', "", va.text)
            school_l.append(school)

    num = 0
    for table in tables:
        school_n = school_l[num]
        # print(school_n)
        # print(len(table.rows))
        for e in table.rows:
            if e.cells[1].text == "姓名":
                continue
            else:
                name = e.cells[1].text
            print(school_n,name)
            try:
                in_datas(school_n, name, year)
            except Exception as e:
                print(e)
        num+=1

    # print(table)

if __name__ == '__main__':
    file = u"D:\桌面\\ruxuan\新建文件夹"
    # create_table()
    read_dir(file)
    # run()
